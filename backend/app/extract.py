"""从上传的文档中抽取纯文本, 供模型分析或润色。

支持: .docx / .pdf / .xlsx / .xls / .csv / .txt / .md
返回: {"ok": bool, "text": str, "kind": str, "truncated": bool} 或 {"ok": False, "error": str}
"""
from __future__ import annotations

import io

import pandas as pd

MAX_CHARS = 50000  # 防止超长文档撑爆上下文/额度


def _cap(text: str) -> tuple[str, bool]:
    text = text.strip()
    if len(text) > MAX_CHARS:
        return text[:MAX_CHARS] + "\n\n…（内容过长已截断）", True
    return text, False


def _table_text(df: pd.DataFrame, max_rows: int = 200) -> str:
    if len(df) > max_rows:
        df = df.head(max_rows)
    return df.to_string(index=False)


def extract_text(filename: str, content: bytes) -> dict:
    name = (filename or "").lower()
    try:
        if name.endswith(".docx"):
            from docx import Document

            doc = Document(io.BytesIO(content))
            parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for t in doc.tables:
                for row in t.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))
            text, truncated = _cap("\n".join(parts))
            return {"ok": True, "text": text, "kind": "docx", "truncated": truncated}

        if name.endswith(".pdf"):
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(content))
            pages = [(page.extract_text() or "") for page in reader.pages]
            joined = "\n\n".join(p.strip() for p in pages if p.strip())
            if not joined:
                return {"ok": False, "error": "这个 PDF 没有可提取的文字（可能是扫描件/图片型 PDF）。"}
            text, truncated = _cap(joined)
            return {"ok": True, "text": text, "kind": "pdf", "truncated": truncated}

        if name.endswith((".xlsx", ".xls")):
            df = pd.read_excel(io.BytesIO(content))
            text, truncated = _cap(_table_text(df))
            return {"ok": True, "text": text, "kind": "excel", "truncated": truncated}

        if name.endswith(".csv"):
            df = pd.read_csv(io.BytesIO(content))
            text, truncated = _cap(_table_text(df))
            return {"ok": True, "text": text, "kind": "csv", "truncated": truncated}

        if name.endswith((".txt", ".md")):
            text, truncated = _cap(content.decode("utf-8", "ignore"))
            return {"ok": True, "text": text, "kind": "text", "truncated": truncated}

        return {"ok": False, "error": "暂不支持这种文件类型（支持 Word/PDF/Excel/CSV/txt）。"}
    except Exception as e:  # noqa: BLE001
        return {"ok": False, "error": f"解析文件失败：{e}"}
