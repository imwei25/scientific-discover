"""伦理审查文书生成 (knowingly programmatic, 不读外部 .docx 模板)。

支持 4 种模板:
  - informed_consent     知情同意书
  - protocol             研究方案
  - crf                  病例报告表(CRF)
  - data_use_commitment  数据使用承诺

设计:
  - 模板用 python-docx 程序化生成: 结构稳定, 易维护, 占位符高亮明显;
  - 占位符约定: {字段名}, 调用时 fields={"研究名称": "...", ...} 替换;
  - 缺失字段保留 [占位] 标记, 而不是空白(让审查者一眼能看到要补充什么)。

接口:
  render(template: str, fields: dict) -> bytes  # 返回 .docx 字节流
"""
from __future__ import annotations

import io
from typing import Callable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


# ---------- 共享样式工具 ----------

_PLACEHOLDER_COLOR = RGBColor(0xD9, 0x77, 0x06)  # 暖色琥珀(D6 设计变量), 让占位醒目


def _add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)


def _add_section(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)


def _add_para(doc: Document, *parts: str | tuple[str, str]) -> None:
    """段落写入: 字符串 = 普通文本; tuple("placeholder","key") = 占位高亮 [key]。"""
    p = doc.add_paragraph()
    for part in parts:
        if isinstance(part, tuple) and len(part) == 2 and part[0] == "placeholder":
            r = p.add_run(f"[{part[1]}]")
            r.font.color.rgb = _PLACEHOLDER_COLOR
            r.bold = True
        else:
            p.add_run(str(part))


def _fill(value: str | None, key: str) -> str | tuple[str, str]:
    """如果 value 非空返回值, 否则返回 ('placeholder', key) 让 _add_para 高亮。"""
    if value is None or str(value).strip() == "":
        return ("placeholder", key)
    return str(value)


# ---------- 4 个模板 ----------

def _render_informed_consent(doc: Document, f: dict) -> None:
    _add_title(doc, "知情同意书")
    doc.add_paragraph(
        "本草案需经伦理委员会(IRB/EC)审核批准后方可使用。"
    ).runs[0].italic = True

    _add_section(doc, "一、研究信息")
    _add_para(doc, "研究名称: ", _fill(f.get("研究名称"), "研究名称"))
    _add_para(doc, "研究者: ", _fill(f.get("研究者"), "研究者"))
    _add_para(doc, "所属机构: ", _fill(f.get("机构"), "机构"))
    _add_para(doc, "联系方式: ", _fill(f.get("联系方式"), "联系方式"))
    _add_para(doc, "日期: ", _fill(f.get("日期"), "日期"))

    _add_section(doc, "二、研究目的")
    _add_para(doc, _fill(f.get("研究目的"), "研究目的"))

    _add_section(doc, "三、研究流程")
    doc.add_paragraph(
        "您将被邀请参与本研究, 大致流程包括: 入组评估、按方案接受相应检测/干预、"
        "随访与数据采集。具体步骤会由研究人员当面说明。"
    )

    _add_section(doc, "四、潜在风险")
    _add_para(doc, _fill(f.get("风险"), "风险"))

    _add_section(doc, "五、可能的获益")
    _add_para(doc, _fill(f.get("受益"), "受益"))

    _add_section(doc, "六、自愿参加与退出")
    doc.add_paragraph(
        "您完全自愿参加本研究, 可在任何时间退出, 不影响您今后接受医疗服务的权利。"
    )

    _add_section(doc, "七、隐私与数据保密")
    doc.add_paragraph(
        "您的个人信息将被严格保密, 数据仅用于本研究目的。发表时不会暴露可识别的个人信息。"
    )

    _add_section(doc, "八、费用与补偿")
    doc.add_paragraph("研究相关检测/干预的费用承担与补偿安排, 请向研究者咨询。")

    _add_section(doc, "九、研究相关损害的处理")
    doc.add_paragraph("如发生与研究相关的健康损害, 研究方将按伦理委员会批准的方案给予处理。")

    _add_section(doc, "十、签字栏")
    doc.add_paragraph("受试者签字: _______________   日期: _______________")
    doc.add_paragraph("法定代理人签字(如适用): _______________   日期: _______________")
    doc.add_paragraph("研究者签字: _______________   日期: _______________")


def _render_protocol(doc: Document, f: dict) -> None:
    _add_title(doc, "研究方案")

    _add_section(doc, "1. 研究基本信息")
    _add_para(doc, "研究名称: ", _fill(f.get("研究名称"), "研究名称"))
    _add_para(doc, "主要研究者(PI): ", _fill(f.get("研究者"), "研究者"))
    _add_para(doc, "承担机构: ", _fill(f.get("机构"), "机构"))
    _add_para(doc, "联系方式: ", _fill(f.get("联系方式"), "联系方式"))
    _add_para(doc, "起止日期: ", _fill(f.get("日期"), "日期"))

    _add_section(doc, "2. 研究背景与目的")
    _add_para(doc, _fill(f.get("研究目的"), "研究目的"))

    _add_section(doc, "3. 研究设计")
    _add_para(doc, _fill(f.get("研究设计"), "研究设计"))

    _add_section(doc, "4. 入选与排除标准")
    _add_para(doc, "入选标准: ", _fill(f.get("入选标准"), "入选标准"))
    _add_para(doc, "排除标准: ", _fill(f.get("排除标准"), "排除标准"))

    _add_section(doc, "5. 样本量")
    _add_para(doc, _fill(f.get("样本量"), "样本量"))

    _add_section(doc, "6. 主要/次要终点")
    _add_para(doc, "主要终点: ", _fill(f.get("主要终点"), "主要终点"))
    _add_para(doc, "次要终点: ", _fill(f.get("次要终点"), "次要终点"))

    _add_section(doc, "7. 干预/操作流程")
    _add_para(doc, _fill(f.get("干预措施"), "干预措施"))

    _add_section(doc, "8. 统计分析计划")
    _add_para(doc, _fill(f.get("统计分析"), "统计分析"))

    _add_section(doc, "9. 风险与获益评估")
    _add_para(doc, "风险: ", _fill(f.get("风险"), "风险"))
    _add_para(doc, "受益: ", _fill(f.get("受益"), "受益"))

    _add_section(doc, "10. 伦理与知情同意")
    doc.add_paragraph(
        "本研究将提交所在机构伦理委员会审查, 所有受试者签署知情同意书后方可入组。"
    )

    _add_section(doc, "11. 数据管理与保密")
    doc.add_paragraph("所有数据将去标识化处理, 严格保密, 仅本研究使用。")


def _render_crf(doc: Document, f: dict) -> None:
    _add_title(doc, "病例报告表 (CRF)")

    _add_para(doc, "研究名称: ", _fill(f.get("研究名称"), "研究名称"))
    _add_para(doc, "受试者编号: ____________   入组日期: ____________")
    doc.add_paragraph()

    _add_section(doc, "一、基本信息")
    table = doc.add_table(rows=4, cols=2)
    table.style = "Light Grid Accent 1"
    cells = [
        ("性别", "□ 男  □ 女"),
        ("出生年份", "______"),
        ("身高 (cm)", "______"),
        ("体重 (kg)", "______"),
    ]
    for i, (k, v) in enumerate(cells):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    _add_section(doc, "二、入选与排除评估")
    _add_para(doc, "入选标准全部满足: □ 是  □ 否")
    _add_para(doc, "排除标准均不满足: □ 是  □ 否")

    _add_section(doc, "三、基线评估")
    doc.add_paragraph("(由研究人员根据方案要求填写)")
    doc.add_paragraph("__________________________________________________")
    doc.add_paragraph("__________________________________________________")

    _add_section(doc, "四、干预/治疗记录")
    _add_para(doc, "干预类型: ", _fill(f.get("干预措施"), "干预措施"))
    doc.add_paragraph("开始日期: ____________   结束日期: ____________")

    _add_section(doc, "五、终点指标")
    _add_para(doc, "主要终点: ", _fill(f.get("主要终点"), "主要终点"))
    _add_para(doc, "次要终点: ", _fill(f.get("次要终点"), "次要终点"))

    _add_section(doc, "六、不良事件 (AE) 记录")
    doc.add_paragraph("□ 无  □ 有(请详述, 含严重程度、持续时间、与研究的相关性)")
    doc.add_paragraph("__________________________________________________")

    _add_section(doc, "七、研究者签字")
    doc.add_paragraph("研究者: _______________   日期: _______________")


def _render_data_use_commitment(doc: Document, f: dict) -> None:
    _add_title(doc, "数据使用承诺书")

    doc.add_paragraph(
        "本承诺书由研究项目主要研究者出具, 用于声明对研究数据使用、存储、共享与"
        "保密的承诺, 提交伦理委员会备案。"
    )

    _add_section(doc, "一、项目信息")
    _add_para(doc, "研究名称: ", _fill(f.get("研究名称"), "研究名称"))
    _add_para(doc, "主要研究者: ", _fill(f.get("研究者"), "研究者"))
    _add_para(doc, "承担机构: ", _fill(f.get("机构"), "机构"))
    _add_para(doc, "联系方式: ", _fill(f.get("联系方式"), "联系方式"))

    _add_section(doc, "二、数据来源与范围")
    _add_para(doc, _fill(f.get("数据来源"), "数据来源"))

    _add_section(doc, "三、使用承诺")
    doc.add_paragraph("本人/本团队承诺:")
    doc.add_paragraph("1. 仅将所获数据用于上述研究目的, 不用于其他任何用途;")
    doc.add_paragraph("2. 对涉及个人隐私的数据进行去标识化处理, 严格保密;")
    doc.add_paragraph("3. 数据存储在受控环境, 访问权限仅授予研究授权人员;")
    doc.add_paragraph("4. 不擅自向第三方提供原始数据; 如需共享, 须报伦理委员会批准;")
    doc.add_paragraph("5. 研究结束后按规定保存数据, 保存期限到期后按规定销毁。")

    _add_section(doc, "四、保存与销毁")
    _add_para(doc, "保存期限: ", _fill(f.get("保存期限"), "保存期限"))
    _add_para(doc, "存储位置: ", _fill(f.get("存储位置"), "存储位置"))

    _add_section(doc, "五、签字")
    doc.add_paragraph("承诺人(研究者): _______________   日期: ", )
    _add_para(doc, "日期: ", _fill(f.get("日期"), "日期"))
    doc.add_paragraph("机构盖章: _______________")


# ---------- 调度 ----------

_TEMPLATES: dict[str, Callable[[Document, dict], None]] = {
    "informed_consent": _render_informed_consent,
    "protocol": _render_protocol,
    "crf": _render_crf,
    "data_use_commitment": _render_data_use_commitment,
}


def list_templates() -> list[str]:
    return list(_TEMPLATES.keys())


def render(template: str, fields: dict | None = None) -> bytes:
    """根据模板与字段渲染 .docx 字节流。

    template ∈ {'informed_consent','protocol','crf','data_use_commitment'}
    fields 缺失项以高亮 [占位] 形式保留。
    """
    builder = _TEMPLATES.get(template)
    if not builder:
        raise ValueError(f"未知伦理材料模板: {template}")
    doc = Document()
    # 全局默认字体(中英文)
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)

    builder(doc, fields or {})

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
