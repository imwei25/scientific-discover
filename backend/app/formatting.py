"""把重排后的文本生成 Word(.docx) 投稿稿。

不依赖 LaTeX/pandoc, 直接用 python-docx 输出 Word, 安装轻、跨平台。
按目标期刊的"投稿稿"版式规格(journals.py 的 docx 字段)程序化设定:
页面/页边距、正文中英文字体与字号、行距、连续行号——这些是投稿稿(单栏)
的核心要求, 效果远好于从零拼装。印刷双栏终稿不在此范围(属出版社流程)。
对常见 Markdown 标记(#, ##, **, - )做基础解析, 映射到 Word 标题/正文样式。
"""
from __future__ import annotations

import io
import re

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from .journals import get_docx_spec, get_journal

_BOLD = re.compile(r"\*\*(.+?)\*\*")
_SECTION_BRACKET = re.compile(r"^【.+】$")
_BULLET = re.compile(r"^[-*]\s+")

# 纸张尺寸(宽 x 高, cm)
_PAGE = {"a4": (21.0, 29.7), "letter": (21.59, 27.94)}
_SPACING_RULE = {
    1.0: WD_LINE_SPACING.SINGLE,
    1.5: WD_LINE_SPACING.ONE_POINT_FIVE,
    2.0: WD_LINE_SPACING.DOUBLE,
}
# sectPr 内 lnNumType 之后允许出现的元素(用于 schema 感知插入, 保证 OOXML 顺序合法)
_SECTPR_AFTER_LNNUM = (
    "w:pgNumType", "w:cols", "w:formProt", "w:vAlign", "w:noEndnote",
    "w:titlePg", "w:textDirection", "w:bidi", "w:rtlGutter", "w:docGrid",
    "w:printerSettings", "w:sectPrChange",
)


def _add_runs_with_bold(paragraph, text: str) -> None:
    """处理一行内的 **加粗** 标记。"""
    pos = 0
    for m in _BOLD.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        run = paragraph.add_run(m.group(1))
        run.bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _apply_page_and_style(doc: Document, spec: dict) -> None:
    """按期刊规格设定页面、页边距、正文中英文字体/字号、行距。"""
    w, h = _PAGE.get(spec["page"], _PAGE["a4"])
    margin = Cm(float(spec["margin_cm"]))
    for section in doc.sections:
        section.page_width = Cm(w)
        section.page_height = Cm(h)
        section.left_margin = section.right_margin = margin
        section.top_margin = section.bottom_margin = margin

    normal = doc.styles["Normal"]
    normal.font.name = spec["body_font"]
    normal.font.size = Pt(float(spec["body_size"]))
    # 中文字体(eastAsia)需写进 rPr 的 rFonts, python-docx 无高层 API。
    cjk = spec.get("body_font_cjk")
    if cjk:
        rpr = normal.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), cjk)
        rfonts.set(qn("w:ascii"), spec["body_font"])
        rfonts.set(qn("w:hAnsi"), spec["body_font"])
    rule = _SPACING_RULE.get(float(spec["line_spacing"]))
    if rule is not None:
        normal.paragraph_format.line_spacing_rule = rule


def _add_line_numbers(doc: Document, count_by: int = 1, start: int = 1,
                      restart: str = "continuous") -> None:
    """往首个 section 的 sectPr 注入连续行号(schema 感知插入, 不用 append)。"""
    sectPr = doc.sections[0]._sectPr
    for el in sectPr.findall(qn("w:lnNumType")):
        sectPr.remove(el)
    ln = OxmlElement("w:lnNumType")
    ln.set(qn("w:countBy"), str(count_by))
    ln.set(qn("w:start"), str(start))
    ln.set(qn("w:restart"), restart)
    # OOXML CT_SectPr 要求 lnNumType 在 pgMar 之后、cols 之前; 用 insert_element_before
    # 保证落在合法位置(直接 append 会排到 cols/docGrid 之后, 严格校验器会丢弃)。
    sectPr.insert_element_before(ln, *_SECTPR_AFTER_LNNUM)


def build_docx(text: str, journal_id: str = "", references: list[str] | None = None) -> bytes:
    doc = Document()
    spec = get_docx_spec(journal_id)
    _apply_page_and_style(doc, spec)

    journal = get_journal(journal_id)
    if journal:
        title = doc.add_heading(journal["name"] + " · 排版稿", level=0)
        title.alignment = 1  # center

    for raw in text.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif _SECTION_BRACKET.match(line.strip()):
            doc.add_heading(line.strip().strip("【】"), level=2)
        elif _BULLET.match(line):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_bold(p, _BULLET.sub("", line))
        else:
            p = doc.add_paragraph()
            _add_runs_with_bold(p, line)

    # 追加按期刊样式格式化好的参考文献
    if references:
        cn = bool(journal) and journal_id == "general_cn"
        doc.add_heading("参考文献" if cn else "References", level=1)
        for ref in references:
            ref = (ref or "").strip()
            if ref:
                doc.add_paragraph(ref)

    # 连续行号放最后(确保 sectPr 已是最终状态)
    if spec.get("line_numbers"):
        _add_line_numbers(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
